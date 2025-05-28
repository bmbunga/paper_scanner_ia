from docx import Document
from io import BytesIO

def generate_word(summary_text: str, title: str = "Résumé généré par Paper Scanner IA", source: str = None) -> BytesIO:
    doc = Document()
    doc.add_heading(title, level=1)
    if source:
        doc.add_paragraph(f"Source : {source}", style='Intense Quote')

    blocks = summary_text.split("\n")
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        if block.lower().startswith("résumé"):
            doc.add_heading("Résumé", level=2)
            doc.add_paragraph(block.split(":", 1)[-1].strip())
        elif block.lower().startswith("moléc"):
            doc.add_heading("Molécules mentionnées", level=2)
            doc.add_paragraph(block.split(":", 1)[-1].strip())
        elif block.lower().startswith("patho"):
            doc.add_heading("Pathologies ciblées", level=2)
            doc.add_paragraph(block.split(":", 1)[-1].strip())
        elif block.lower().startswith("type"):
            doc.add_heading("Type d'étude", level=2)
            doc.add_paragraph(block.split(":", 1)[-1].strip())
        elif block.lower().startswith("auteur"):
            doc.add_heading("Auteurs principaux", level=2)
            doc.add_paragraph(block.split(":", 1)[-1].strip())
        else:
            doc.add_paragraph(block)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
